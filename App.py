import os
import datetime
import base64
import nltk
import spacy
import streamlit as st
import pandas as pd
from docx import Document
from pyresparser import ResumeParser

nltk.download('punkt')
nltk.download('stopwords')
spacy.load('en_core_web_sm')

# Function to extract text from PDF
def pdf_reader(file_path):
    import PyPDF2
    with open(file_path, 'rb') as file:
        pdf = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf.pages:
            text += page.extract_text()
    return text

# Fallback parser if ResumeParser fails
def fallback_resume_data(resume_text):
    from nltk.tokenize import word_tokenize
    words = word_tokenize(resume_text)
    data = {
        'name': "Name Placeholder",
        'email': next((w for w in words if '@' in w and '.' in w), "Email Placeholder")
    }
    return data

# Streamlit App
def run():
    st.set_page_config(page_title="Smart Resume Analyzer", layout="wide")
    st.title("Smart Resume Analyzer :sunglasses:")

    menu = ["User", "Admin"]
    choice = st.sidebar.selectbox("Select Role", menu)

    if choice == "User":
        st.subheader("Upload Your Resume")
        uploaded_file = st.file_uploader("Upload Resume", type=["pdf", "docx", "doc"])

        if uploaded_file is not None:
            save_path = os.path.join("Uploaded_Resumes", uploaded_file.name)
            with open(save_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
            st.success("Resume uploaded successfully")

            # Extract text
            if uploaded_file.name.endswith(".pdf"):
                resume_text = pdf_reader(save_path)
            elif uploaded_file.name.endswith((".doc", ".docx")):
                doc = Document(save_path)
                resume_text = "\n".join([para.text for para in doc.paragraphs])
            else:
                resume_text = ""

            # Parse data
            try:
                resume_data = ResumeParser(save_path).get_extracted_data()
                if not resume_data or not resume_data.get("email"):
                    resume_data = fallback_resume_data(resume_text)
            except Exception:
                resume_data = fallback_resume_data(resume_text)

            st.markdown("### Extracted Data")
            st.write(f"**Name:** {resume_data.get('name', 'N/A')}")
            st.write(f"**Email:** {resume_data.get('email', 'N/A')}")

    else:
        st.subheader("Admin Login")
        ad_user = st.text_input("Username")
        ad_password = st.text_input("Password", type='password')

        if st.button("Login"):
            if ad_user == 'nikhil' and ad_password == 'Nik@123':
                st.success("Welcome Group 6")

                folder_path = "./Uploaded_Resumes/"
                if not os.path.exists(folder_path):
                    os.makedirs(folder_path)

                resume_files = [
                    f for f in os.listdir(folder_path)
                    if f.endswith((".pdf", ".docx", ".doc"))
                ]

                data = []
                for idx, file in enumerate(resume_files, start=1):
                    file_path = os.path.join(folder_path, file)

                    if file.endswith(".pdf"):
                        resume_text = pdf_reader(file_path)
                    elif file.endswith((".doc", ".docx")):
                        doc = Document(file_path)
                        resume_text = "\n".join([para.text for para in doc.paragraphs])
                    else:
                        resume_text = ""

                    try:
                        resume_data = ResumeParser(file_path).get_extracted_data()
                        if not resume_data or not resume_data.get("email"):
                            resume_data = fallback_resume_data(resume_text)
                    except Exception:
                        resume_data = fallback_resume_data(resume_text)

                    name = resume_data.get("name", "Name Placeholder")
                    email = resume_data.get("email", "Email Placeholder")
                    timestamp = datetime.datetime.fromtimestamp(os.path.getmtime(file_path))

                    data.append({
                        "ID": idx,
                        "Name": name,
                        "Email": email,
                        "Timestamp": timestamp.strftime("%Y-%m-%d %H:%M:%S")
                    })

                df = pd.DataFrame(data)
                st.markdown("### User's Resume Data")
                st.dataframe(df, use_container_width=True)

                csv = df.to_csv(index=False).encode('utf-8')
                st.download_button(
                    label="Download Report",
                    data=csv,
                    file_name="resume_analysis_report.csv",
                    mime="text/csv"
                )
            else:
                st.error("Wrong ID & Password Provided")

if __name__ == '__main__':
    run()
